# -*- coding: utf-8 -*-
"""SpectroMind Audit Raporu docx oluşturucu"""
import os

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx gerekli: pip install python-docx")
    exit(1)

def create_audit_docx():
    doc = Document()
    
    # Başlık
    title = doc.add_heading('SpectroMind / SpectroMaster Audit Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Omniscient Prime v31.0 — Bilimsel Doğruluk ve Algoritma Denetimi')
    doc.add_paragraph('Tarih: 2025-02-04 | Kapsam: Toplu SpectroMind çıktısı (40+ molekül)')
    doc.add_paragraph()
    
    # 1) Executive Summary
    doc.add_heading('1) Executive Summary', level=1)
    items = [
        "0.x İntegralli Protonlar: Simülasyonda değişebilir (OH/NH) protonlar exchInteg = 0.1 + rnd() * 0.5 ile üretiliyor; QC'de yanlış parse ediliyor.",
        "Kritik Kod Hatası: h1AromaticIntegral hesaplamasında parseInt+replace kullanımı; '0.547' yanlış parse ediliyor.",
        "Sistem Yanlış FAIL Verdi: Benzen, THF, DMSO gibi moleküller geçerli veri ile FAIL aldı.",
        "VETO_H1_AROMATIC_INTEGRAL Aşırı Sert: Parse hatası veto tetikliyor.",
        "FP_RULES Kapsamı: Çoğu molekül için fingerprint yok.",
        "IR Ar C=C Bandı: 1450–1600 cm⁻¹ bandı neredeyse hiç 'gözlenmedi'.",
        "HSQC Atama Oranı: Birçok molekülde <%75.",
        "calculateTotalIntegral Mantığı: Broad peak için v>=0.5 ? v : 1 tutarsız.",
        "MS ppm Hatası: 13C izotop peak'leri yüksek ppm hata.",
        "Öncelik: P0=aromatik integral parse; P1=exchangeable tutarlılık; P2=IR threshold; P3=FP_RULES."
    ]
    for i, it in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {it}', style='List Number')
    
    # 2) Kanonik Veri Özeti
    doc.add_heading('2) Kanonik Veri Özeti', level=1)
    doc.add_paragraph('Kimlik: context, formula, smiles, exact_mass, DBE. Spektral: 1H, 13C, MS, IR, 2D. Sistem: coverage, PASS/FAIL, data_completeness.')
    
    # 3) Detaylı Doğruluk Tablosu
    doc.add_heading('3) Detaylı Doğruluk Tablosu', level=1)
    table = doc.add_table(rows=10, cols=8)
    table.style = 'Table Grid'
    headers = ['Modül', 'Sistem', 'Benim', 'Çelişki', 'Kanıt', 'Kök Sebep', 'Öncelik', 'Patch']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['Formula/DBE', 'PASS', 'PASS', 'NONE', '-', '-', '-', '-'],
        ['1H', 'FAIL/WARN', 'WARN', 'OVERSTRICT', '0.x parse', 'CODE_BUG', 'P0', 'parseFloat'],
        ['13C', 'WARN', 'PASS', 'UNDERSTRICT', 'Simetri', 'RULE_DESIGN', 'P3', 'Simetri faktörü'],
        ['HSQC', 'WARN', 'WARN', 'NONE', '-', '-', '-', '-'],
        ['COSY', 'WARN', 'WARN', 'NONE', '-', '-', '-', '-'],
        ['HMBC', 'PASS', 'PASS', 'NONE', '-', '-', '-', '-'],
        ['MS', 'PASS/FAIL', 'PASS', 'FALSE_NEG', '13C adduct', 'THRESHOLD', 'P2', 'Exclude 13C'],
        ['IR', 'WARN', 'WARN', 'OVERSTRICT', 'Ar C=C', 'THRESHOLD', 'P2', 'Band threshold'],
        ['Library', 'Partial', 'Partial', 'NONE', 'FP_RULES', 'LIBRARY_GAP', 'P3', 'Genişlet']
    ]
    for r, row in enumerate(data, 1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val
    
    # 4) Kök Sebep
    doc.add_heading('4) Kök Sebep Analizi', level=1)
    doc.add_heading('5.1 CODE_BUG', level=2)
    doc.add_paragraph('h1AromaticIntegral: parseInt(String(s.integ).replace(/\\D/g, ""), 10) ondalık değerleri bozuyor. Çözüm: parseFloat veya parseInteg.')
    doc.add_heading('5.2 LOGIC_BUG', level=2)
    doc.add_paragraph('calculateTotalIntegral broadContrib: (v>=0.5)?v:1 tutarsız. Exchangeable için tutarlı politika.')
    doc.add_heading('5.3 RULE_DESIGN', level=2)
    doc.add_paragraph('VETO_H1_AROMATIC_INTEGRAL eşik sert. aromatic_integral_tolerance_ratio ile soft geçiş.')
    
    # 5) 0.x İntegral Analizi
    doc.add_heading('5) 0.x İntegralli Protonlar Analizi', level=1)
    doc.add_paragraph('Tanım: 0.x değerler değişebilir (OH/NH) protonları temsil eder. Kaynak: exchInteg = 0.1 + rnd()*0.5 simülasyon.')
    doc.add_paragraph('Karar: Gerçek sinyal Evet, Raporlamada Evet, QC veto None. Aksiyon: Exchangeable say, integral check hariç.')
    
    # 6) Patch Planı
    doc.add_heading('6) Patch Planı', level=1)
    patches = [
        ('P0', '1H', 'h1AromaticIntegral parseFloat', '~16490'),
        ('P1', '1H', 'calculateTotalIntegral broadContrib', '~15541'),
        ('P1', 'Simülasyon', 'exchInteg 0.5-1.0', '~4079'),
        ('P2', 'IR', 'Ar C=C band threshold', 'verifyIR'),
        ('P3', 'FP_RULES', 'Eksik SMILES ekle', 'FP_RULES array')
    ]
    for p0, mod, act, loc in patches:
        doc.add_paragraph(f'{p0}: {mod} — {act} ({loc})', style='List Bullet')
    
    # 7) Regression Test
    doc.add_heading('7) Regression Test Planı', level=1)
    doc.add_paragraph('10 test: Fenolik OH, simetrik aromatik, solvent, low SNR HSQC, 13C overlap, MS [M+Na]+, IR band, fingerprint yok/var, integral impurity.')
    
    # 8) Self-Audit
    doc.add_heading('8) Self-Audit & Evrim', level=1)
    doc.add_paragraph('Kesinlik için: D2O exchange, DEPT-135, uzun delay 13C. Kural motoru: classifyZeroishIntegral, aromatic_integral_parse_fix. Changelog: P0-P3 patch listesi.')
    
    # Kaydet
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SPECTROMIND_AUDIT_RAPORU.docx')
    doc.save(out_path)
    print(f'Kaydedildi: {out_path}')
    return out_path

if __name__ == '__main__':
    create_audit_docx()
